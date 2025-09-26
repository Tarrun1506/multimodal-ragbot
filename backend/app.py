from flask import Flask, request, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
import uuid
from datetime import datetime
import numpy as np
from sentence_transformers import SentenceTransformer
from pymongo import MongoClient
import PyPDF2
import docx
import requests
import faiss
import whisper
# Add these imports at the top
import base64
import wave
import tempfile
import os
from speech_service import speech_service

# Try new import first, fallback to old
try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ImportError:
    from langchain.text_splitter import RecursiveCharacterTextSplitter

# Document class for text processing
class Document:
    def __init__(self, page_content, metadata=None):
        self.page_content = page_content
        self.metadata = metadata or {}
# Image processing imports
from PIL import Image
import pytesseract
import cv2
import easyocr
import io
import base64

app = Flask(__name__)
CORS(app)

# Configuration
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'mp3', 'wav', 'm4a', 'ogg'}
OLLAMA_BASE_URL = "http://localhost:11434"
ALLOWED_MODELS = ['gemma3:1b', 'mistral:latest', 'llama3.2:1b']

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# MongoDB setup
try:
    client = MongoClient('mongodb://localhost:27017/', serverSelectionTimeoutMS=5000)
    client.admin.command('ping')
    db = client['starrag_bot']
    documents_collection = db['documents']
    conversations_collection = db['conversations']
    print("✅ MongoDB connected successfully")
except Exception as e:
    print(f"⚠️ MongoDB connection failed: {e}")
    documents_collection = None
    conversations_collection = None
    client = None

embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
faiss_index = None
document_chunks = []

# Initialize EasyOCR reader once at startup
print("🔧 Initializing EasyOCR...")
try:
    # Fix PIL.Image.ANTIALIAS deprecation issue
    import PIL.Image
    if not hasattr(PIL.Image, 'ANTIALIAS'):
        PIL.Image.ANTIALIAS = PIL.Image.LANCZOS
    
    easyocr_reader = easyocr.Reader(['en'])
    print("✅ EasyOCR initialized successfully")
except Exception as e:
    print(f"⚠️ EasyOCR initialization failed: {e}")
    easyocr_reader = None

# Initialize Whisper model for audio transcription
print("🔧 Loading Whisper model for audio transcription...")
try:
    whisper_model = whisper.load_model("base")  # You can use "small", "medium", "large" for better accuracy
    print("✅ Whisper model loaded successfully")
except Exception as e:
    print(f"⚠️ Failed to load Whisper model: {e}")
    whisper_model = None

class RAGPipeline:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        self.embedding_model = embedding_model
        
    def allowed_file(self, filename):
        return '.' in filename and \
               filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

    def extract_text_from_image(self, image_path):
        """Extract text from images using OCR"""
        global easyocr_reader
        
        try:
            # Method 1: Try EasyOCR first (if available)
            if easyocr_reader is not None:
                try:
                    results = easyocr_reader.readtext(image_path)
                    if results:
                        # Extract text from EasyOCR results
                        text = ' '.join([result[1] for result in results])
                        if text.strip():
                            print(f"✅ EasyOCR extracted text: {len(text)} characters")
                            return text
                except Exception as e:
                    print(f"⚠️ EasyOCR failed, falling back to Tesseract: {e}")
            
            # Method 2: Fallback to Tesseract OCR (if available)
            print("🔧 Attempting Tesseract OCR...")
            try:
                # Check if Tesseract is available
                pytesseract.get_tesseract_version()
                
                # Preprocess image for better OCR
                opencv_image = cv2.imread(image_path)
                if opencv_image is None:
                    # Try with PIL if OpenCV fails
                    image = Image.open(image_path)
                    text = pytesseract.image_to_string(image, config='--psm 6')
                else:
                    gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
                    
                    # Apply image preprocessing for better OCR
                    # Noise removal
                    denoised = cv2.medianBlur(gray, 5)
                    
                    # Thresholding
                    _, thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                    
                    # Convert back to PIL Image
                    processed_image = Image.fromarray(thresh)
                    
                    # Extract text using Tesseract
                    text = pytesseract.image_to_string(processed_image, config='--psm 6')
                
                if text.strip():
                    print(f"✅ Tesseract extracted text: {len(text)} characters")
                    return text.strip()
                else:
                    return "No text found in image"
                    
            except Exception as tesseract_error:
                print(f"⚠️ Tesseract not available: {tesseract_error}")
                # If both OCR methods fail, return a basic message
                return "Image uploaded successfully, but text extraction failed. Please install Tesseract OCR for better results."
            
        except Exception as e:
            print(f"❌ Error extracting text from image {image_path}: {e}")
            # Return a user-friendly message instead of failing completely
            return "Image uploaded successfully, but text extraction encountered an error. The image has been processed and stored."
    
    def extract_text_from_file(self, file_path, filename):
        """Extract text from different file types including images"""
        text = ""
        file_extension = filename.rsplit('.', 1)[1].lower()
        
        try:
            if file_extension == 'txt':
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
            elif file_extension == 'pdf':
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            elif file_extension == 'docx':
                doc = docx.Document(file_path)
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
            elif file_extension in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff']:
                text = self.extract_text_from_image(file_path)
                # For images, we always return some text even if OCR fails
                if not text or text.startswith('Error'):
                    text = f"Image file: {filename} (OCR text extraction was attempted)"
        except Exception as e:
            print(f"Error extracting text from {filename}: {str(e)}")
            return None
            
        return text

    def process_document(self, file_path, filename):
        """Process a document and add it to the vector store"""
        global faiss_index, document_chunks
        
        # Extract text
        text = self.extract_text_from_file(file_path, filename)
        if not text:
            return False
            
        # Create document and split into chunks
        doc = Document(page_content=text, metadata={"source": filename})
        chunks = self.text_splitter.split_documents([doc])
        
        # Generate embeddings
        chunk_texts = [chunk.page_content for chunk in chunks]
        embeddings = self.embedding_model.encode(chunk_texts)
        
        # Store in MongoDB
        doc_id = str(uuid.uuid4())
        document_data = {
            "_id": doc_id,
            "filename": filename,
            "content": text,
            "chunks": [
                {
                    "text": chunk.page_content,
                    "metadata": chunk.metadata,
                    "embedding": embedding.tolist()
                }
                for chunk, embedding in zip(chunks, embeddings)
            ],
            "created_at": datetime.utcnow()
        }
        
        if documents_collection is not None:
            try:
                documents_collection.insert_one(document_data)
            except Exception as e:
                print(f"Warning: Could not store document in MongoDB: {e}")
        else:
            print("Warning: MongoDB not available - document not persisted")
        
        # Add to FAISS index
        if faiss_index is None:
            dimension = embeddings.shape[1]
            faiss_index = faiss.IndexFlatL2(dimension)
            document_chunks = []
        
        faiss_index.add(embeddings.astype('float32'))
        
        for i, chunk in enumerate(chunks):
            document_chunks.append({
                "text": chunk.page_content,
                "metadata": chunk.metadata,
                "doc_id": doc_id
            })
        
        return True

    def similarity_search(self, query, k=5):
        """Search for similar chunks using FAISS"""
        global faiss_index, document_chunks
        
        if faiss_index is None or len(document_chunks) == 0:
            return []
        
        # Generate query embedding
        query_embedding = self.embedding_model.encode([query])
        
        # Search in FAISS
        distances, indices = faiss_index.search(query_embedding.astype('float32'), k)
        
        # Return relevant chunks
        results = []
        for i, idx in enumerate(indices[0]):
            if idx < len(document_chunks):
                chunk = document_chunks[idx]
                results.append({
                    "text": chunk["text"],
                    "metadata": chunk["metadata"],
                    "score": float(distances[0][i])
                })
        
        return results
    
    def format_response(self, raw_response):
        """Format the response with proper markdown and mathematical notation"""
        import re
        
        # Clean up the response
        formatted = raw_response.strip()
        
        # Fix bullet points - handle both line-start and inline bullets
        # First, handle bullets at the start of lines
        formatted = re.sub(r'^\*\s+([^\n]+)', r'- \1', formatted, flags=re.MULTILINE)
        formatted = re.sub(r'^-\s+([^\n]+)', r'- \1', formatted, flags=re.MULTILINE)
        
        # Handle inline bullet lists - only if there are multiple bullets
        lines = formatted.split('\n')
        new_lines = []
        
        for line in lines:
            if '•' in line and line.count('•') > 1:
                # Only process lines with multiple bullets (inline lists)
                parts = re.split(r'\s*•\s*', line)
                # Process each part
                for part in parts:
                    clean_part = part.strip()
                    if clean_part:  # Skip empty parts
                        new_lines.append(f'- {clean_part}')
            elif '•' in line:
                # Single bullet - just convert to markdown
                clean_line = line.replace('•', '-').strip()
                new_lines.append(clean_line)
            else:
                new_lines.append(line)
        
        formatted = '\n'.join(new_lines)
        
        # Clean up formatting issues
        formatted = re.sub(r'^\s*-\s*$', '', formatted, flags=re.MULTILINE)  # Remove empty bullets
        # Don't remove all double newlines - preserve paragraph breaks for regular text
        formatted = re.sub(r'\n\s*\n\s*\n+', '\n\n', formatted)  # Only remove triple+ newlines
        
        # Keep long bullet points intact - don't break them
        # The user wants complete information, not truncated descriptions
        
        # Fix nested bullets
        formatted = re.sub(r'^\s+\*\s+([^\n]+)', r'  - \1', formatted, flags=re.MULTILINE)
        formatted = re.sub(r'^\s+-\s+([^\n]+)', r'  - \1', formatted, flags=re.MULTILINE)
        
        # Fix numbered lists
        formatted = re.sub(r'(\d+)\.\s*([^\n]+)', r'\1. \2', formatted)
        
        # Fix bold formatting - ensure proper markdown
        formatted = re.sub(r'\*\*([^*]+)\*\*', r'**\1**', formatted)
        
        # Fix mathematical formulas - convert to proper LaTeX format (more conservative)
        # Only apply math formatting if the response seems to contain mathematical content
        if any(char in formatted for char in ['=', '+', '-', '*', '/', '%']) and len(formatted.split()) < 50:
            # Handle simple equations
            formatted = re.sub(r'\b([A-Za-z_]+)\s*=\s*([A-Za-z0-9_+\-*/\s()]+)(?=\s|$)', r'$\1 = \2$', formatted)
            # Handle fractions in equations
            formatted = re.sub(r'\b(\d+)/(\d+)\b', r'$\\frac{\1}{\2}$', formatted)
        # Handle percentages
        formatted = re.sub(r'(\d+(?:\.\d+)?)\s*%', r'\1%', formatted)
        
        # Fix headers - ensure proper spacing (more conservative)
        formatted = re.sub(r'^\*\*([A-Z][A-Za-z\s]{3,}):?\*\*\s*$', r'## \1', formatted, flags=re.MULTILINE)
        # Don't convert single words or short phrases to headers
        
        # Ensure proper line breaks for readability
        formatted = re.sub(r'\n\s*\n\s*\n+', '\n\n', formatted)
        # Final cleanup for bullet formatting
        # Remove any remaining bullet symbols that weren't converted
        formatted = re.sub(r'•', '', formatted)
        # Clean up multiple consecutive newlines but preserve paragraph structure
        formatted = re.sub(r'\n{4,}', '\n\n', formatted)
        # Only remove empty lines between bullets if they are actually bullet lists
        if '- ' in formatted and formatted.count('- ') > 2:
            formatted = re.sub(r'\n\s*\n(-\s)', r'\n\1', formatted)
        # Ensure proper spacing before bullet lists (but not between bullets)
        formatted = re.sub(r'([^\n-])\n(-\s)', r'\1\n\n\2', formatted)
        
        # Fix common formatting issues
        formatted = re.sub(r'\*\s*\*([^*]+)\*\s*\*', r'**\1**', formatted)
        
        return formatted.strip()
    
    def generate_response(self, query, context_chunks, model="gemma3:1b"):
        """Generate response using Ollama with proper formatting"""
        # Filter context for maximum relevance to the specific query
        query_lower = query.lower()
        
        # For specific factual queries, prioritize chunks with the answer
        if any(word in query_lower for word in ['who is', 'what is', 'when is', 'where is', 'which is']):
            # Sort chunks by relevance to the specific question
            relevant_chunks = []
            for chunk in context_chunks:
                chunk_lower = chunk["text"].lower()
                relevance_score = 0
                
                # Check for direct answers to the question type
                if 'publisher' in query_lower and 'publish' in chunk_lower:
                    relevance_score += 10
                elif 'author' in query_lower and 'author' in chunk_lower:
                    relevance_score += 10
                elif 'date' in query_lower and any(word in chunk_lower for word in ['date', 'year', 'published']):
                    relevance_score += 10
                
                # Add chunks with high relevance first
                if relevance_score > 5:
                    relevant_chunks.insert(0, chunk)
                else:
                    relevant_chunks.append(chunk)
            
            # Use only the most relevant chunks (max 2 for specific questions)
            context_chunks = relevant_chunks[:2]
        
        context = "\n\n".join([chunk["text"] for chunk in context_chunks])
        
        prompt = f"""You are a precise AI assistant. Answer the specific question using ONLY the most relevant information from the context.

Context:
{context}

Question: {query}

Instructions:
1. Answer ONLY what is specifically asked - be direct and concise
2. If the question asks for a specific fact, provide that fact clearly
3. Don't include additional information unless directly relevant to the question
4. Use **bold text** for the key answer
5. For simple answers, use plain text. For lists with multiple items, use bullet points:
   - Item 1
   - Item 2
   - Item 3
6. For mathematical formulas, use proper notation
7. If the context doesn't contain the specific information asked, say so clearly
8. Keep your response natural and readable

Provide a focused, direct answer:"""

        try:
            # Check Ollama health
            health_check = requests.get(f"{OLLAMA_BASE_URL}/api/tags", timeout=30)
            if health_check.status_code != 200:
                return "Error: Ollama is not responding. Please run 'ollama serve'."
            
            # Verify model is available
            models = health_check.json().get('models', [])
            model_names = [m.get('name', '') for m in models]
            
            if model not in model_names:
                # Try to find an available allowed model
                available_models = [m for m in ALLOWED_MODELS if m in model_names]
                if available_models:
                    model = available_models[0]
                    print(f"Using available model: {model}")
                else:
                    available = ', '.join(model_names) if model_names else 'none'
                    return f"Error: None of the allowed models are available. Available models: {available}"
            
            # Generate response
            response = requests.post(
                f"{OLLAMA_BASE_URL}/api/generate",
                json={
                    "model": model,
                    "prompt": prompt,
                    "stream": False,
                    "options": {
                        "temperature": 0.7,
                        "top_p": 0.9,
                        "num_predict": 2048
                    }
                },
                timeout=300
            )
            
            if response.status_code == 200:
                raw_response = response.json().get("response", "Sorry, I couldn't generate a response.")
                return self.format_response(raw_response)
            else:
                error_msg = f"Ollama error (status {response.status_code})"
                try:
                    error_detail = response.json().get('error', '')
                    if error_detail:
                        error_msg += f": {error_detail}"
                except:
                    pass
                return f"Error: {error_msg}"
                
        except requests.exceptions.ConnectionError:
            return "Error: Cannot connect to Ollama. Please run 'ollama serve'."
        except requests.exceptions.Timeout:
            return "Error: Request timed out. Please try again."
        except Exception as e:
            return f"Error: {str(e)}"

# Initialize RAG pipeline
rag_pipeline = RAGPipeline()

# API Endpoints
@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'message': 'StarRAG Bot API is running',
        'version': '1.0'
    })

@app.route('/api/upload', methods=['POST'])
def upload_file():
    """Handle file upload with better error handling including audio files"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if file and rag_pipeline.allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_path = os.path.join(UPLOAD_FOLDER, filename)
            file_ext = filename.rsplit('.', 1)[1].lower()
            
            # Save file with error handling
            try:
                file.save(file_path)
                print(f"📁 Processing file: {filename}")
            except Exception as save_error:
                print(f"❌ Error saving file: {save_error}")
                return jsonify({'error': f'Failed to save file: {str(save_error)}'}), 500

            # Handle audio files with Whisper transcription
            if file_ext in ['mp3', 'wav', 'm4a', 'ogg']:
                if whisper_model is None:
                    return jsonify({'error': 'Whisper model not loaded. Audio transcription unavailable.'}), 500
                
                try:
                    print(f"🎵 Transcribing audio file: {filename}")
                    result = whisper_model.transcribe(file_path)
                    transcription = result.get('text', '').strip()
                    
                    if transcription:
                        print(f"✅ Transcription completed: {len(transcription)} characters")
                        
                        # Create a text document from the transcription
                        doc_filename = f"{filename}_transcript.txt"
                        doc_path = os.path.join(UPLOAD_FOLDER, doc_filename)
                        
                        with open(doc_path, 'w', encoding='utf-8') as f:
                            f.write(f"Audio Transcription from {filename}:\n\n{transcription}")
                        
                        # Process the transcription as a document
                        success = rag_pipeline.process_document(doc_path, doc_filename)
                        
                        # Clean up temporary transcript file
                        if os.path.exists(doc_path):
                            os.remove(doc_path)
                        
                        if success:
                            return jsonify({
                                'message': f'Audio {filename} transcribed and processed successfully',
                                'transcription': transcription[:500] + ('...' if len(transcription) > 500 else ''),
                                'full_length': len(transcription)
                            })
                        else:
                            return jsonify({'error': 'Failed to process transcription'}), 500
                    else:
                        return jsonify({'error': 'No speech detected in audio file'}), 400
                        
                except Exception as transcription_error:
                    print(f"❌ Error transcribing audio: {transcription_error}")
                    return jsonify({'error': f'Failed to transcribe audio: {str(transcription_error)}'}), 500
                finally:
                    # Clean up original audio file
                    if os.path.exists(file_path):
                        os.remove(file_path)
            
            # Handle regular document files
            else:
                try:
                    success = rag_pipeline.process_document(file_path, filename)
                except Exception as process_error:
                    print(f"❌ Error processing document: {process_error}")
                    # Clean up file even if processing fails
                    if os.path.exists(file_path):
                        os.remove(file_path)
                    return jsonify({'error': f'Failed to process document: {str(process_error)}'}), 500
                
                # Clean up file
                try:
                    if os.path.exists(file_path):
                        os.remove(file_path)
                except Exception as cleanup_error:
                    print(f"⚠️ Warning: Failed to clean up file: {cleanup_error}")
                
                if success:
                    print(f"✅ Successfully processed: {filename}")
                    return jsonify({'message': f'File {filename} processed successfully'})
                else:
                    return jsonify({'error': 'Failed to process document'}), 500
        
        return jsonify({'error': 'Invalid file type. Supported: TXT, PDF, DOCX, Images (PNG, JPG, etc.), Audio (MP3, WAV, M4A, OGG)'}), 400
        
    except Exception as e:
        print(f"❌ Unexpected error in upload_file: {e}")
        return jsonify({'error': f'Server error: {str(e)}'}), 500

@app.route('/api/documents', methods=['GET'])
def get_documents():
    """Get list of uploaded documents"""
    if documents_collection is None:
        return jsonify({
            'documents': [],
            'message': 'MongoDB not available'
        })
    
    try:
        docs = list(documents_collection.find({}, {'filename': 1, 'created_at': 1}))
        return jsonify({
            'documents': [
                {
                    'id': str(doc['_id']),
                    'filename': doc['filename'],
                    'created_at': doc['created_at'].isoformat()
                }
                for doc in docs
            ]
        })
    except Exception as e:
        return jsonify({
            'documents': [],
            'error': str(e)
        })

@app.route('/api/documents/<doc_id>', methods=['DELETE'])
def delete_document(doc_id):
    """Delete a document"""
    if documents_collection is None:
        return jsonify({'error': 'MongoDB not available'}), 500
    
    try:
        # Remove document from MongoDB
        result = documents_collection.delete_one({'_id': doc_id})
        
        if result.deleted_count == 0:
            return jsonify({'error': 'Document not found'}), 404
            
        # Rebuild FAISS index without this document
        global faiss_index, document_chunks
        
        if faiss_index is not None:
            all_docs = list(documents_collection.find())
            
            faiss_index = None
            document_chunks = []
            
            if all_docs:
                all_embeddings = []
                
                for doc in all_docs:
                    for chunk in doc.get("chunks", []):
                        embedding = np.array(chunk["embedding"])
                        all_embeddings.append(embedding)
                        document_chunks.append({
                            "text": chunk["text"],
                            "metadata": chunk["metadata"],
                            "doc_id": doc["_id"]
                        })
                
                if all_embeddings:
                    embeddings_array = np.vstack(all_embeddings)
                    dimension = embeddings_array.shape[1]
                    faiss_index = faiss.IndexFlatL2(dimension)
                    faiss_index.add(embeddings_array.astype('float32'))
        
        return jsonify({'message': 'Document deleted successfully'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/models', methods=['GET'])
def get_models():
    """Get available Ollama models"""
    try:
        response = requests.get(f"{OLLAMA_BASE_URL}/api/tags", timeout=15)
        if response.status_code == 200:
            models = response.json().get('models', [])
            model_names = [model['name'] for model in models]
            
            # Filter to only show allowed models that are available
            available_models = [m for m in ALLOWED_MODELS if m in model_names]
            
            if available_models:
                return jsonify({
                    'models': available_models,
                    'status': 'connected'
                })
            else:
                return jsonify({
                    'models': ALLOWED_MODELS,
                    'status': 'no_models',
                    'message': f'None of the allowed models found. Available models: {model_names}'
                })
        else:
            return jsonify({
                'models': ALLOWED_MODELS,
                'status': 'error',
                'message': f'Ollama responded with status {response.status_code}'
            })
    except requests.exceptions.ConnectionError:
        return jsonify({
            'models': ALLOWED_MODELS,
            'status': 'disconnected',
            'message': 'Cannot connect to Ollama. Run "ollama serve" in terminal.'
        })
    except Exception as e:
        return jsonify({
            'models': ALLOWED_MODELS,
            'status': 'error',
            'message': str(e)
        })

@app.route('/api/query', methods=['POST'])
def query():
    """Handle user queries"""
    data = request.json
    user_query = data.get('query', '')
    model = data.get('model', 'gemma3:1b')
    conversation_id = data.get('conversation_id')
    
    if not user_query:
        return jsonify({'error': 'No query provided'}), 400
    
    # Validate model
    if model not in ALLOWED_MODELS:
        return jsonify({'error': f'Model {model} is not allowed'}), 400
    
    # Search for relevant chunks
    relevant_chunks = rag_pipeline.similarity_search(user_query, k=5)
    
    if not relevant_chunks:
        return jsonify({
            'response': 'I have no relevant information. Please upload documents first.',
            'sources': [],
            'conversation_id': conversation_id
        })
    
    # Generate response
    response = rag_pipeline.generate_response(user_query, relevant_chunks, model)
    
    # Create or update conversation
    if not conversation_id:
        conversation_id = str(uuid.uuid4())
    
    conversation_data = {
        'conversation_id': conversation_id,
        'query': user_query,
        'response': response,
        'model': model,
        'sources': list(set([chunk['metadata'].get('source', 'Unknown') for chunk in relevant_chunks])),
        'timestamp': datetime.utcnow()
    }
    
    if conversations_collection is not None:
        try:
            conversations_collection.insert_one(conversation_data)
        except Exception as e:
            print(f"Warning: Could not store conversation: {e}")
    
    return jsonify({
        'response': response,
        'sources': list(set([chunk['metadata'].get('source', 'Unknown') for chunk in relevant_chunks])),
        'conversation_id': conversation_id
    })

@app.route('/api/conversations', methods=['GET'])
def list_conversations():
    """List all conversations"""
    if conversations_collection is None:
        return jsonify({'conversations': [], 'message': 'MongoDB not available'})
    
    try:
        pipeline = [
            {"$sort": {"timestamp": 1}},
            {"$group": {
                "_id": "$conversation_id",
                "first_query": {"$first": "$query"},
                "first_timestamp": {"$first": "$timestamp"},
                "model": {"$first": "$model"}
            }},
            {"$sort": {"first_timestamp": -1}}
        ]
        conversations = list(conversations_collection.aggregate(pipeline))
        return jsonify({
            'conversations': [
                {
                    'conversation_id': c['_id'],
                    'first_query': c.get('first_query', ''),
                    'first_timestamp': c.get('first_timestamp').isoformat() if c.get('first_timestamp') else '',
                    'model': c.get('model', '')
                }
                for c in conversations
            ]
        })
    except Exception as e:
        return jsonify({'conversations': [], 'error': str(e)})

@app.route('/api/conversations/<conversation_id>', methods=['GET'])
def get_conversation(conversation_id):
    """Get conversation history"""
    if conversations_collection is None:
        return jsonify({'conversations': [], 'message': 'MongoDB not available'})
    
    try:
        conversations = list(conversations_collection.find(
            {'conversation_id': conversation_id},
            {'_id': 0}
        ).sort('timestamp', 1))
        
        return jsonify({'conversations': conversations})
    except Exception as e:
        return jsonify({'conversations': [], 'error': str(e)})

if __name__ == '__main__':
    # Load existing documents on startup
    if documents_collection is not None:
        try:
            docs = list(documents_collection.find())
            if docs:
                all_embeddings = []
                document_chunks = []
                
                for doc in docs:
                    for chunk in doc.get("chunks", []):
                        embedding = np.array(chunk["embedding"])
                        all_embeddings.append(embedding)
                        document_chunks.append({
                            "text": chunk["text"],
                            "metadata": chunk["metadata"],
                            "doc_id": doc["_id"]
                        })
                
                if all_embeddings:
                    embeddings_array = np.vstack(all_embeddings)
                    dimension = embeddings_array.shape[1]
                    faiss_index = faiss.IndexFlatL2(dimension)
                    faiss_index.add(embeddings_array.astype('float32'))
                    print(f"✅ Loaded {len(docs)} documents with {len(document_chunks)} chunks")
        except Exception as e:
            print(f"⚠️ Error loading existing documents: {e}")

@app.route('/api/voice/speech-to-text', methods=['POST'])
def speech_to_text():
    """Convert speech audio to text"""
    try:
        data = request.json
        
        if not data or 'audio_data' not in data:
            return jsonify({'error': 'No audio data provided'}), 400
        
        # Extract audio data and format
        audio_base64 = data['audio_data']
        audio_format = data.get('format', 'webm')  # Default to webm for browser recordings
        language = data.get('language', 'en-US')
        
        # Decode base64 audio data
        try:
            audio_bytes = base64.b64decode(audio_base64.split(',')[-1] if ',' in audio_base64 else audio_base64)
        except Exception as e:
            return jsonify({'error': f'Invalid audio data: {str(e)}'}), 400
        
        # Convert speech to text
        text = speech_service.speech_to_text(audio_bytes, audio_format, language)
        
        if text:
            return jsonify({
                'text': text,
                'confidence': 'high',
                'language': language
            })
        else:
            return jsonify({
                'error': 'Could not transcribe audio',
                'text': '',
                'confidence': 'low'
            }), 400
            
    except Exception as e:
        print(f"❌ Speech-to-text error: {e}")
        return jsonify({'error': f'Speech recognition failed: {str(e)}'}), 500

@app.route('/api/voice/text-to-speech', methods=['POST'])
def text_to_speech():
    """Convert text to speech (optional - for voice responses)"""
    try:
        data = request.json
        
        if not data or 'text' not in data:
            return jsonify({'error': 'No text provided'}), 400
        
        text = data['text']
        # For now, return a message about TTS availability
        # You can integrate with external TTS services like Google TTS, Azure TTS, etc.
        
        return jsonify({
            'message': 'Text-to-speech functionality can be integrated with external services',
            'suggestions': [
                'Google Cloud Text-to-Speech',
                'Amazon Polly',
                'Azure Cognitive Services',
                'OpenAI TTS'
            ]
        })
        
    except Exception as e:
        return jsonify({'error': f'Text-to-speech error: {str(e)}'}), 500

@app.route('/api/voice/supported-languages', methods=['GET'])
def get_supported_languages():
    """Get list of supported languages for speech recognition"""
    languages = {
        'en-US': 'English (US)',
        'en-GB': 'English (UK)',
        'es-ES': 'Spanish',
        'fr-FR': 'French',
        'de-DE': 'German',
        'it-IT': 'Italian',
        'pt-BR': 'Portuguese (Brazil)',
        'ru-RU': 'Russian',
        'ja-JP': 'Japanese',
        'ko-KR': 'Korean',
        'zh-CN': 'Chinese (Simplified)',
        'hi-IN': 'Hindi'
    }
    return jsonify({'languages': languages})

@app.route('/api/voice/whisper-status', methods=['GET'])
def whisper_status():
    """Check if Whisper model is loaded and ready"""
    return jsonify({
        'whisper_loaded': whisper_model is not None,
        'model_type': 'base' if whisper_model is not None else None
    })

@app.route('/api/voice/transcribe-audio', methods=['POST'])
def transcribe_audio():
    """Transcribe audio data using Whisper"""
    try:
        if whisper_model is None:
            return jsonify({'error': 'Whisper model not loaded. Audio transcription unavailable.'}), 500
        
        # Get audio data from request
        data = request.get_json()
        if not data or 'audio' not in data:
            return jsonify({'error': 'No audio data provided'}), 400
        
        audio_base64 = data['audio']
        
        # Decode base64 audio data
        try:
            audio_data = base64.b64decode(audio_base64)
            print(f"📊 Audio data decoded: {len(audio_data)} bytes")
        except Exception as e:
            print(f"❌ Base64 decode error: {e}")
            return jsonify({'error': f'Invalid base64 audio data: {str(e)}'}), 400
        
        # Create temporary file for audio (use webm extension since that's what we're receiving)
        with tempfile.NamedTemporaryFile(suffix='.webm', delete=False) as temp_audio:
            temp_audio.write(audio_data)
            temp_audio.flush()
            
            try:
                # Transcribe using Whisper (Whisper can handle WebM format)
                print(f"Transcribing recorded audio...")
                result = whisper_model.transcribe(temp_audio.name)
                transcription = result.get('text', '').strip()
                
                if transcription:
                    print(f"Audio transcription completed: {len(transcription)} characters")
                    return jsonify({
                        'success': True,
                        'transcription': transcription,
                        'length': len(transcription)
                    })
                else:
                    return jsonify({'error': 'No speech detected in audio'}), 400
                    
            except Exception as transcription_error:
                print(f"Error transcribing audio: {transcription_error}")
                return jsonify({'error': f'Failed to transcribe audio: {str(transcription_error)}'}), 500
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_audio.name)
                except:
                    pass
                    
    except Exception as e:
        print(f"Unexpected error in transcribe_audio: {e}")
        return jsonify({'error': f'Server error: {str(e)}'}), 500

if __name__ == '__main__':
    print("🚀 Starting StarRAG Bot server...")
    print("📡 Server will be available at: http://localhost:5000")
    print("🎯 API endpoints ready for transcription and document processing")
    app.run(debug=True, host='0.0.0.0', port=5000)